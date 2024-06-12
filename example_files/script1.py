import asyncio
import logging
import os
from docx import Document
from ai_engine_sdk import (
    AiEngine,
    is_agent_message,
    is_ai_engine_message,
    is_confirmation_message,
    is_stop_message,
    is_task_selection_message, TaskSelectionMessage,
    ApiBaseMessage, FunctionGroup
)

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

api_key = os.getenv("AV_API_KEY", "")
file_path = 'IBP_Problemstatement.docx'  # Path to the uploaded document
output_path = 'Extracted_Data.docx'  # Path to the new document to be created

async def extract_data():
    doc = Document(file_path)
    extracted_data = []

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 3:  # Ensuring there are at least three columns for Term, Status, and Definition
                term = cells[0].text.strip()
                status = cells[1].text.strip()
                definition = cells[2].text.strip()
                if term and status and definition:
                    extracted_data.append((term, status, definition))
    return extracted_data

async def create_new_document(data):
    new_doc = Document()
    new_doc.add_heading('Extracted Data', level=1)

    for term, status, definition in data:
        new_doc.add_heading(term, level=2)
        new_doc.add_paragraph(f"Status: {status}")
        new_doc.add_paragraph(f"Definition: {definition}")
        new_doc.add_paragraph("\n")

    new_doc.save(output_path)
    logger.info(f"Data successfully extracted and saved to {output_path}")

async def main():
    logger.debug("🚀 Starting data extraction and document creation process")

    ai_engine = AiEngine(api_key)

    # Extract data from the document
    data = await extract_data()
    if data:
        logger.debug("Data extracted successfully")
        await create_new_document(data)
    else:
        logger.error("Failed to extract data")

if __name__ == "__main__":
    asyncio.run(main())
