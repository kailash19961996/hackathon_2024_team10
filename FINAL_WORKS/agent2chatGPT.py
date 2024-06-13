import os
import json
import logging
import openai
from docx import Document
from uagents import Agent, Context, Model
from uagents.setup import fund_agent_if_low

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# OpenAI API key
api_key = os.getenv('OPENAI_API_KEY')
if api_key is None:
    raise ValueError("API key not found. Please set the OPENAI_API_KEY environment variable.")
openai.api_key = api_key
print("---------api_key_works---------")

# Define model for claim
class Claim(Model):
    claim: str

# Agent2 setup
agent2 = Agent(name="agent2", port=8008, seed="agent2 secret phrase", endpoint=["http://127.0.0.1:8008/submit"])
print("uAgent address: ", agent2.address)  # This needs to be used for the recipient_address in agent3
# print("Fetch network address: ", agent2.wallet.address())

# Fund the agent if low
fund_agent_if_low(agent2.wallet.address())

# File paths
file_path = 'documentB_claim_details.docx'  # Path to the input document
text_output_path = 'documentB_claims.txt'  # Path to the text file to be created

def extract_text_from_document(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return '\n'.join(full_text)

def extract_rules_with_chatgpt(document_text):
    prompt = f"Extract the following information from the document:\n\n{document_text}\n\n" \
             "Fields to extract:\n" \
             "user_id, name, request, age, ill_health, current_pension, join_date, spouse_eligible" \
             "Put all the information of each user in one line with their corresponding field names seperated by commas"
    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ],
        max_tokens=2000,
        temperature=0.5
    )
    return response.choices[0].message['content'].strip()


@agent2.on_event("startup")
async def create_rules(ctx: Context):
    
    # Extract text from the document
    doc_text = extract_text_from_document(file_path)
    
    # Use ChatGPT to extract rules
    extracted_rules = extract_rules_with_chatgpt(doc_text)
    
    if extracted_rules:
        logger.debug("Rules extracted successfully")
        
        # Save rules to text file
        with open(text_output_path, "w") as f:
            f.write(extracted_rules)
        
        print("---------text_file_created---------")
        
        # Send rules to agent3
        recipient_address = "agent1qweavyt3rh8qs0yjytyks7wydlcxfvshzpewt9zsle7txqa22gawytel2kd"
        await ctx.send(recipient_address, Claim(claim=extracted_rules))
        logger.info(f"Rules successfully extracted and sent to agent3")
    else:
        logger.error(" ******** Failed to extract rules from the document ******** ")

if __name__ == "__main__":
    agent2.run()
