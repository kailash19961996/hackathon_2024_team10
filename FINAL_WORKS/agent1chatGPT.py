import os
import logging
import openai
from docx import Document
from uagents import Agent, Context, Model

logging.basicConfig(level=logging.DEBUG) # Configure logging
logger = logging.getLogger(__name__)

api_key = os.getenv('OPENAI_API_KEY') # OpenAI API key
if api_key is None:
    raise ValueError("API key not found. Please set the OPENAI_API_KEY environment variable.")
openai.api_key = api_key
print("---------api_key_works---------")

class Rule(Model): # Define model for Rule
    rules: str

# Agent1 setup
agent1 = Agent(name="agent1", port=8010, seed="agent1 secret phrase", endpoint=["http://127.0.0.1:80/submit"])
print("uAgent address: ", agent1.address)

# File paths
file_path = 'documentA_insurance_terms.docx'  # Path to the input document
text_output_path = 'documentA_rules.txt'  # Path to the text file to be created

# 
def extract_text_from_document(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return '\n'.join(full_text)

def extract_rules_with_chatgpt(doc_text):
    prompt = f"Extract the important rules from the following insurance agreement text and add a approximate price range for each rule at the end of rule:\n\n{doc_text}"
    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ],
        max_tokens=2000,
        temperature=0.5
    )
    return response.choices[0].message['content'].strip()

@agent1.on_event("startup")
async def create_rules(ctx: Context):
    logger.debug("🚀 Starting rule extraction from document")
    
    # Extract text from the document
    doc_text = extract_text_from_document(file_path)
    
    # Use ChatGPT to extract rules
    extracted_rules = extract_rules_with_chatgpt(doc_text)
    
    if extracted_rules:
        logger.debug("Rules extracted successfully")
        
        # Save rules to text file
        with open(text_output_path, "w") as f:
            f.write(extracted_rules)
        
        print("---------text_file_created---------")
        
        # Send rules to agent3
        recipient_address = "agent1qweavyt3rh8qs0yjytyks7wydlcxfvshzpewt9zsle7txqa22gawytel2kd"
        await ctx.send(recipient_address, Rule(rules=extracted_rules))
        logger.info(f"Rules successfully extracted and sent to agent3")
    else:
        logger.error(" ******** Failed to extract rules from the document ******** ")

if __name__ == "__main__":
    agent1.run()

